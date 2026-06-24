# Copyright 2026 sharexpress
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND.
#

import io
import boto3
import asyncio
import logging
from uuid import uuid4
from datetime import datetime
from typing import List, Dict, Any, Optional
from botocore.client import Config
from fastapi import HTTPException

from docx import Document
from docx.shared import Pt

from core.config import (
    MINIO_ACCESS_KEY,
    MINIO_SECRET_KEY,
    MINIO_REGION,
    MINIO_BUCKET,
    MINIO_ENDPOINT_INTERNAL,
    MINIO_ENDPOINT_PUBLIC,
)
from core.database import get_db

logger = logging.getLogger(__name__)
db = get_db()

s3_internal = boto3.client(
    "s3",
    endpoint_url=MINIO_ENDPOINT_INTERNAL,
    aws_access_key_id=MINIO_ACCESS_KEY,
    aws_secret_access_key=MINIO_SECRET_KEY,
    region_name=MINIO_REGION,
    config=Config(
        signature_version="s3v4",
        s3={"addressing_style": "path"},
        connect_timeout=5,
        read_timeout=15,
        retries={"max_attempts": 2},
    ),
)

s3_public = boto3.client(
    "s3",
    endpoint_url=MINIO_ENDPOINT_PUBLIC,
    aws_access_key_id=MINIO_ACCESS_KEY,
    aws_secret_access_key=MINIO_SECRET_KEY,
    region_name=MINIO_REGION,
    config=Config(
        signature_version="s3v4",
        s3={"addressing_style": "path"},
    ),
)

SUPPORTED_TYPES = {"pdf", "docx", "doc", "txt", "md", "html", "json", "csv"}
SUPPORTED_TEXT_TYPES = {"txt", "md", "html", "json", "csv"}


async def _get_file(file_id: str, session: dict):
    """Shared helper — fetch file doc from DB with session auth check."""
    file_doc = await db.files.find_one(
        {
            "file_id": file_id,
            "is_deleted": False,
        }
    )
    if not file_doc:
        raise HTTPException(status_code=404, detail="File not found")

    # Access control: check if the file belongs to this active sharing session
    if str(file_doc["sharing_session_id"]) != str(session["sharing_session_ID"]):
        raise HTTPException(
            status_code=403,
            detail="Access denied: file does not belong to this sharing session",
        )

    return file_doc


class EditorController:
    @staticmethod
    async def load_file(file_id: str, session: dict):
        """Return metadata + presigned URL."""
        file_doc = await _get_file(file_id, session)
        storage_key = file_doc["storage_key"]
        filename = file_doc["original_name"]
        ext = filename.rsplit(".", 1)[-1].lower()

        if ext not in SUPPORTED_TYPES:
            raise HTTPException(
                status_code=400, detail=f"Unsupported file type: .{ext}"
            )

        url = await asyncio.to_thread(
            s3_public.generate_presigned_url,
            "get_object",
            Params={
                "Bucket": MINIO_BUCKET,
                "Key": storage_key,
            },
            ExpiresIn=3600,
        )

        return {
            "file_id": file_id,
            "filename": filename,
            "ext": ext,
            "url": url,
            "storage_key": storage_key,
        }

    @staticmethod
    async def get_text_content(file_id: str, session: dict):
        """Fetch plain-text file content directly from S3 storage."""
        file_doc = await _get_file(file_id, session)
        filename = file_doc["original_name"]
        ext = filename.rsplit(".", 1)[-1].lower()

        if ext not in SUPPORTED_TEXT_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"File extension .{ext} is not supported for text editing",
            )

        try:
            obj = await asyncio.to_thread(
                s3_internal.get_object,
                Bucket=MINIO_BUCKET,
                Key=file_doc["storage_key"],
            )
            data = await asyncio.to_thread(obj["Body"].read)
            content = data.decode("utf-8")
        except Exception as e:
            logger.error(f"Failed to fetch text content from S3: {e}")
            raise HTTPException(
                status_code=500, detail=f"Failed to read file from storage: {e}"
            )

        return {
            "file_id": file_id,
            "filename": filename,
            "content": content,
            "ext": ext,
        }

    @staticmethod
    async def save_text_content(file_id: str, content: str, session: dict):
        """Save plain-text content back to S3, versioning in file_versions collection."""
        # 1. Verify edit permission flag
        if not session.get("can_edit", False):
            raise HTTPException(
                status_code=403,
                detail="Access denied: Edit permission is disabled in this session",
            )

        file_doc = await _get_file(file_id, session)
        filename = file_doc["original_name"]
        ext = filename.rsplit(".", 1)[-1].lower()

        if ext not in SUPPORTED_TEXT_TYPES:
            raise HTTPException(
                status_code=400, detail="Unsupported text file category"
            )

        try:
            # 2. Get next version number
            version_count = await db.file_versions.count_documents(
                {"file_ID": file_id}
            )

            # If it's the very first edit, create Version 1 from current state
            if version_count == 0:
                v1_key = f"{session['sharing_session_ID']}/{file_id}_v1_{filename}"
                await asyncio.to_thread(
                    s3_internal.copy_object,
                    Bucket=MINIO_BUCKET,
                    CopySource={"Bucket": MINIO_BUCKET, "Key": file_doc["storage_key"]},
                    Key=v1_key,
                )
                await db.file_versions.insert_one(
                    {
                        "version_ID": str(uuid4()),
                        "file_ID": file_id,
                        "version_No": 1,
                        "storage_key": v1_key,
                        "editedBy": "session",
                        "editedByUserID": session["sender_ID"],
                        "createdAt": datetime.utcnow(),
                    }
                )
                version_count = 1

            next_version = version_count + 1
            new_version_key = f"{session['sharing_session_ID']}/{file_id}_v{next_version}_{filename}"

            # 3. Save new content as next version file in S3
            content_bytes = content.encode("utf-8")
            await asyncio.to_thread(
                s3_internal.put_object,
                Bucket=MINIO_BUCKET,
                Key=new_version_key,
                Body=content_bytes,
                ContentType=file_doc["mime_type"],
            )

            # 4. Overwrite main file in S3
            await asyncio.to_thread(
                s3_internal.put_object,
                Bucket=MINIO_BUCKET,
                Key=file_doc["storage_key"],
                Body=content_bytes,
                ContentType=file_doc["mime_type"],
            )

            # 5. Insert version metadata
            await db.file_versions.insert_one(
                {
                    "version_ID": str(uuid4()),
                    "file_ID": file_id,
                    "version_No": next_version,
                    "storage_key": new_version_key,
                    "editedBy": "session",
                    "editedByUserID": session["sender_ID"],
                    "createdAt": datetime.utcnow(),
                }
            )

            # Update size in files metadata
            await db.files.update_one(
                {"file_id": file_id},
                {
                    "$set": {
                        "size": len(content_bytes),
                        "updated_at": datetime.utcnow(),
                    }
                },
            )

            # 6. Keep max 10 versions pruning limit (EDIT-04)
            versions = await db.file_versions.find({"file_ID": file_id}).sort(
                "version_No", 1
            ).to_list(length=100)
            if len(versions) > 10:
                prune_count = len(versions) - 10
                for i in range(prune_count):
                    v_to_prune = versions[i]
                    # Delete from S3
                    try:
                        await asyncio.to_thread(
                            s3_internal.delete_object,
                            Bucket=MINIO_BUCKET,
                            Key=v_to_prune["storage_key"],
                        )
                    except Exception as e:
                        logger.error(
                            f"Failed to delete pruned version from S3: {e}"
                        )
                    # Delete from DB
                    await db.file_versions.delete_one(
                        {"version_ID": v_to_prune["version_ID"]}
                    )

            return {
                "success": True,
                "message": "File saved successfully",
                "version_No": next_version,
            }

        except Exception as e:
            logger.error(f"Error saving text file content: {e}", exc_info=True)
            raise HTTPException(
                status_code=500, detail=f"Failed to save content: {str(e)}"
            )

    @staticmethod
    async def get_version_history(file_id: str, session: dict):
        """List version history of the document."""
        await _get_file(file_id, session)  # auth verification
        versions = await db.file_versions.find({"file_ID": file_id}).sort(
            "version_No", -1
        ).to_list(length=100)

        # Serialize list
        serialized_versions = []
        for v in versions:
            serialized_versions.append(
                {
                    "version_ID": v["version_ID"],
                    "file_ID": v["file_ID"],
                    "version_No": v["version_No"],
                    "editedBy": v["editedBy"],
                    "editedByUserID": v["editedByUserID"],
                    "createdAt": v["createdAt"].isoformat(),
                }
            )
        return {"file_id": file_id, "versions": serialized_versions}

    @staticmethod
    async def rollback_to_version(
        file_id: str, version_id: str, session: dict
    ):
        """Rollback file state to a previous version."""
        if not session.get("can_edit", False):
            raise HTTPException(
                status_code=403, detail="Access denied: Edit permission is disabled"
            )

        file_doc = await _get_file(file_id, session)
        version_doc = await db.file_versions.find_one(
            {"version_ID": version_id, "file_ID": file_id}
        )

        if not version_doc:
            raise HTTPException(status_code=404, detail="Version not found")

        try:
            # 1. Copy version object back to primary key in S3
            await asyncio.to_thread(
                s3_internal.copy_object,
                Bucket=MINIO_BUCKET,
                CopySource={
                    "Bucket": MINIO_BUCKET,
                    "Key": version_doc["storage_key"],
                },
                Key=file_doc["storage_key"],
            )

            # Get version count for listing
            version_count = await db.file_versions.count_documents(
                {"file_ID": file_id}
            )
            next_version = version_count + 1
            new_version_key = f"{session['sharing_session_ID']}/{file_id}_v{next_version}_{file_doc['original_name']}"

            # 2. Duplicate as next version for logging rollback
            await asyncio.to_thread(
                s3_internal.copy_object,
                Bucket=MINIO_BUCKET,
                CopySource={
                    "Bucket": MINIO_BUCKET,
                    "Key": version_doc["storage_key"],
                },
                Key=new_version_key,
            )

            await db.file_versions.insert_one(
                {
                    "version_ID": str(uuid4()),
                    "file_ID": file_id,
                    "version_No": next_version,
                    "storage_key": new_version_key,
                    "editedBy": "session",
                    "editedByUserID": session["sender_ID"],
                    "createdAt": datetime.utcnow(),
                }
            )

            # Update size in database
            obj_meta = await asyncio.to_thread(
                s3_internal.head_object,
                Bucket=MINIO_BUCKET,
                Key=file_doc["storage_key"],
            )
            await db.files.update_one(
                {"file_id": file_id},
                {
                    "$set": {
                        "size": obj_meta["ContentLength"],
                        "updated_at": datetime.utcnow(),
                    }
                },
            )

            return {
                "success": True,
                "message": f"Successfully rolled back to version {version_doc['version_No']}",
            }

        except Exception as e:
            logger.error(f"Error during rollback: {e}")
            raise HTTPException(status_code=500, detail=f"Rollback failed: {e}")

    @staticmethod
    async def get_docx_content(file_id: str, session: dict):
        """Fetch DOCX from MinIO → parse → return paragraphs as JSON."""
        file_doc = await _get_file(file_id, session)  # auth check
        storage_key = file_doc["storage_key"]

        try:
            obj = await asyncio.to_thread(
                s3_internal.get_object, Bucket=MINIO_BUCKET, Key=storage_key
            )
            data = await asyncio.to_thread(obj["Body"].read)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Storage fetch failed: {e}"
            )

        try:
            doc = Document(io.BytesIO(data))
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to parse DOCX: {e}"
            )

        paragraphs = []
        for para in doc.paragraphs:
            bold, italic, font_size = False, False, 12
            if para.runs:
                r = para.runs[0]
                bold = bool(r.bold)
                italic = bool(r.italic)
                try:
                    font_size = int(r.font.size.pt) if r.font.size else 12
                except Exception:
                    font_size = 12
            paragraphs.append(
                {
                    "text": para.text,
                    "bold": bold,
                    "italic": italic,
                    "font_size": font_size,
                }
            )

        return {"file_id": file_id, "paragraphs": paragraphs}

    @staticmethod
    async def save_docx(file_id: str, content: list, session: dict):
        """Rebuild DOCX from edited paragraphs → overwrite in MinIO."""
        # Verify edit permission flag
        if not session.get("can_edit", False):
            raise HTTPException(
                status_code=403,
                detail="Access denied: Edit permission is disabled in this session",
            )

        file_doc = await _get_file(file_id, session)
        storage_key = file_doc["storage_key"]

        if not content:
            raise HTTPException(status_code=400, detail="No content provided")

        try:
            doc = Document()
            for para in content:
                p = doc.add_paragraph()
                run = p.add_run(para.get("text", ""))
                run.bold = para.get("bold", False)
                run.italic = para.get("italic", False)
                size = para.get("font_size", 12)
                try:
                    run.font.size = Pt(int(size))
                except Exception:
                    run.font.size = Pt(12)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Failed to build DOCX: {e}"
            )

        try:
            await asyncio.to_thread(
                s3_internal.put_object,
                Bucket=MINIO_BUCKET,
                Key=storage_key,
                Body=buf.getvalue(),
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Storage save failed: {e}"
            )

        return {"success": True, "file_id": file_id}
